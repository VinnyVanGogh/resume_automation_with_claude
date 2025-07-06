"""
DOCX generator for resume output.

This module provides the DOCXGenerator class for creating professional,
ATS-compliant Word documents using python-docx with YAML-based template system.
"""

import logging
import yaml
from pathlib import Path
from typing import Optional, Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

from src.models import ResumeData, Experience, Education, Project, Certification, Skills
from .config import DOCXConfig, OutputConfig


logger = logging.getLogger(__name__)


class DOCXGenerator:
    """
    Generate DOCX resume output from structured resume data.
    
    This class uses python-docx to create professional Word documents
    with ATS-compliant formatting and YAML-based template system.
    """
    
    def __init__(self, config: Optional[DOCXConfig] = None, template_name: str = "professional") -> None:
        """
        Initialize DOCX generator with configuration and template.
        
        Args:
            config: DOCX generation configuration. If None, uses defaults.
            template_name: Name of the template to use (professional, modern, minimal, tech)
        """
        self.config = config or DOCXConfig()
        self.template_name = template_name
        self.styles_config = self._load_styles_config()
        self.template_config = self._load_template_config()
        logger.info(f"DOCXGenerator initialized with template: {template_name}")
    
    def _load_styles_config(self) -> Dict[str, Any]:
        """
        Load DOCX styles configuration from YAML file.
        
        Returns:
            Dict containing style configuration
        """
        styles_path = Path(__file__).parent.parent / "templates" / "docx" / "styles.yaml"
        try:
            with open(styles_path, 'r') as f:
                return yaml.safe_load(f)
        except FileNotFoundError:
            logger.warning(f"Styles config not found at {styles_path}, using defaults")
            return {}
        except yaml.YAMLError as e:
            logger.error(f"Error loading styles config: {e}")
            return {}
    
    def _load_template_config(self) -> Dict[str, Any]:
        """
        Load DOCX template configuration from YAML file.
        
        Returns:
            Dict containing template configuration
        """
        template_path = Path(__file__).parent.parent / "templates" / "docx" / "templates.yaml"
        try:
            with open(template_path, 'r') as f:
                return yaml.safe_load(f)
        except FileNotFoundError:
            logger.warning(f"Template config not found at {template_path}, using defaults")
            return {}
        except yaml.YAMLError as e:
            logger.error(f"Error loading template config: {e}")
            return {}
    
    def generate(
        self, 
        resume_data: ResumeData, 
        output_path: Optional[Path] = None,
        template_vars: Optional[Dict[str, Any]] = None
    ) -> bytes:
        """
        Generate DOCX resume from structured resume data.
        
        Args:
            resume_data: Validated resume data from Phase 3 formatter
            output_path: Optional path to save DOCX file
            template_vars: Additional template variables (unused for DOCX)
            
        Returns:
            bytes: Generated DOCX content as bytes
            
        Raises:
            ValueError: If resume data is invalid
            RuntimeError: If DOCX generation fails
        """
        logger.info("Starting DOCX generation")
        
        # Validate input data
        if not isinstance(resume_data, ResumeData):
            raise ValueError("resume_data must be a ResumeData instance")
        
        try:
            # Create new document
            doc = Document()
            
            # Configure document properties
            self._set_document_properties(doc, resume_data)
            
            # Configure document styling
            self._configure_document_styles(doc)
            
            # Set margins
            self._set_page_margins(doc)
            
            # Add content sections
            self._add_header_section(doc, resume_data)
            
            if resume_data.summary:
                self._add_summary_section(doc, resume_data.summary)
            
            if resume_data.experience:
                self._add_experience_section(doc, resume_data.experience)
            
            if resume_data.education:
                self._add_education_section(doc, resume_data.education)
            
            if resume_data.skills:
                self._add_skills_section(doc, resume_data.skills)
            
            if resume_data.projects:
                self._add_projects_section(doc, resume_data.projects)
            
            if resume_data.certifications:
                self._add_certifications_section(doc, resume_data.certifications)
            
            # Save to file if path provided
            if output_path:
                output_path.parent.mkdir(parents=True, exist_ok=True)
                doc.save(output_path)
                logger.info(f"DOCX resume saved to: {output_path}")
            
            # Return as bytes for consistency with other generators
            from io import BytesIO
            doc_bytes = BytesIO()
            doc.save(doc_bytes)
            doc_content = doc_bytes.getvalue()
            
            logger.info("DOCX generation completed successfully")
            return doc_content
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            raise RuntimeError(f"DOCX generation failed: {e}") from e
    
    def _set_document_properties(self, doc: Document, resume_data: ResumeData) -> None:
        """
        Set document properties and metadata.
        
        Args:
            doc: Document object to configure
            resume_data: Resume data for metadata
        """
        core_props = doc.core_properties
        core_props.title = f"{resume_data.contact.name} - Resume"
        core_props.author = resume_data.contact.name
        core_props.subject = "Professional Resume"
        core_props.keywords = "resume, cv, professional"
        core_props.comments = "Generated by Resume Automation Tool"
    
    def _configure_document_styles(self, doc: Document) -> None:
        """
        Configure document-wide styles using YAML configuration.
        
        Args:
            doc: Document object to configure
        """
        styles = doc.styles
        
        # Get style configuration for current template theme
        paragraph_styles = self.styles_config.get('paragraph_styles', {})
        theme_config = self.styles_config.get('themes', {}).get(self.template_name, {})
        
        # Configure Normal style
        self._configure_style(styles['Normal'], paragraph_styles.get('Normal', {}), theme_config)
        
        # Configure all heading styles
        for style_name in ['Heading1', 'Heading2', 'Heading3']:
            if style_name in paragraph_styles:
                self._configure_style(styles[style_name.replace('Heading', 'Heading ')], 
                                    paragraph_styles[style_name], theme_config)
        
        # Configure other paragraph styles
        self._configure_additional_styles(styles, paragraph_styles, theme_config)
    
    def _configure_style(self, style, style_config: Dict[str, Any], theme_config: Dict[str, Any]) -> None:
        """
        Configure a specific style using YAML configuration.
        
        Args:
            style: The style object to configure
            style_config: Style configuration from YAML
            theme_config: Theme-specific overrides
        """
        if not style_config:
            return
            
        # Apply font settings
        font = style.font
        if 'font_name' in style_config:
            font_name = theme_config.get('fonts', {}).get('primary', style_config['font_name'])
            font.name = font_name
        if 'font_size' in style_config:
            font.size = Pt(style_config['font_size'])
        if 'bold' in style_config:
            font.bold = style_config['bold']
        if 'italic' in style_config:
            font.italic = style_config['italic']
            
        # Apply paragraph formatting
        paragraph_format = style.paragraph_format
        if 'line_spacing' in style_config:
            paragraph_format.line_spacing = style_config['line_spacing']
        if 'space_before' in style_config:
            paragraph_format.space_before = Pt(style_config['space_before'])
        if 'space_after' in style_config:
            paragraph_format.space_after = Pt(style_config['space_after'])
        if 'alignment' in style_config:
            alignment_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            paragraph_format.alignment = alignment_map.get(style_config['alignment'], WD_ALIGN_PARAGRAPH.LEFT)
        if 'left_indent' in style_config:
            paragraph_format.left_indent = Inches(style_config['left_indent'])
    
    def _configure_additional_styles(self, styles, paragraph_styles: Dict[str, Any], theme_config: Dict[str, Any]) -> None:
        """
        Configure additional custom styles.
        
        Args:
            styles: Document styles collection
            paragraph_styles: Paragraph style configuration
            theme_config: Theme-specific configuration
        """
        # Add custom styles that might not exist
        custom_styles = ['BulletList', 'ContactInfo', 'DateLocation']
        for style_name in custom_styles:
            if style_name in paragraph_styles:
                try:
                    # Try to get existing style, or create new one
                    try:
                        style = styles[style_name]
                    except KeyError:
                        # Create new style based on Normal
                        style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                    
                    self._configure_style(style, paragraph_styles[style_name], theme_config)
                except Exception as e:
                    logger.warning(f"Could not configure style {style_name}: {e}")
    
    def _set_page_margins(self, doc: Document) -> None:
        """
        Set page margins according to YAML configuration.
        
        Args:
            doc: Document object to configure
        """
        sections = doc.sections
        document_config = self.styles_config.get('document', {})
        margins = document_config.get('margins', {})
        
        for section in sections:
            section.top_margin = Inches(margins.get('top', self.config.margin_top))
            section.bottom_margin = Inches(margins.get('bottom', self.config.margin_bottom))
            section.left_margin = Inches(margins.get('left', self.config.margin_left))
            section.right_margin = Inches(margins.get('right', self.config.margin_right))
    
    def _add_header_section(self, doc: Document, resume_data: ResumeData) -> None:
        """
        Add header section with contact information.
        
        Args:
            doc: Document object
            resume_data: Resume data
        """
        contact = resume_data.contact
        
        # Name as title
        name_para = doc.add_heading(contact.name, level=1)
        
        # Contact information
        contact_lines = []
        
        # Email and phone
        email_phone = []
        if contact.email:
            email_phone.append(str(contact.email))
        if contact.phone:
            email_phone.append(contact.phone)
        if email_phone:
            contact_lines.append(" | ".join(email_phone))
        
        # URLs
        url_line = []
        if contact.linkedin:
            url_line.append(str(contact.linkedin))
        if contact.github:
            url_line.append(str(contact.github))
        if contact.website:
            url_line.append(str(contact.website))
        if url_line:
            contact_lines.append(" | ".join(url_line))
        
        # Location
        if contact.location:
            contact_lines.append(contact.location)
        
        # Add contact lines
        for line in contact_lines:
            para = doc.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_summary_section(self, doc: Document, summary: str) -> None:
        """
        Add professional summary section.
        
        Args:
            doc: Document object
            summary: Summary text
        """
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(summary)
    
    def _add_experience_section(self, doc: Document, experiences: List[Experience]) -> None:
        """
        Add experience section.
        
        Args:
            doc: Document object
            experiences: List of experience items
        """
        doc.add_heading("Experience", level=2)
        
        for exp in experiences:
            # Job title and company
            title_para = doc.add_heading(f"{exp.title} | {exp.company}", level=3)
            
            # Dates and location
            date_location = []
            if exp.start_date:
                end_date = exp.end_date or "Present"
                date_location.append(f"{exp.start_date} - {end_date}")
            if exp.location:
                date_location.append(exp.location)
            
            if date_location:
                date_para = doc.add_paragraph(" | ".join(date_location))
                date_run = date_para.runs[0]
                date_run.italic = True
            
            # Responsibilities (bullets field in Experience model)
            if exp.bullets:
                for responsibility in exp.bullets:
                    # Use bullet points (• character)
                    bullet_para = doc.add_paragraph(f"• {responsibility}")
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
    
    def _add_education_section(self, doc: Document, education: List[Education]) -> None:
        """
        Add education section.
        
        Args:
            doc: Document object
            education: List of education items
        """
        doc.add_heading("Education", level=2)
        
        for edu in education:
            # Degree and school
            title_para = doc.add_heading(f"{edu.degree} | {edu.school}", level=3)
            
            # Dates and location
            date_info = []
            if edu.start_date:
                end_date = edu.end_date or "Present"
                date_info.append(f"{edu.start_date} - {end_date}")
            if edu.location:
                date_info.append(edu.location)
            
            if date_info:
                date_para = doc.add_paragraph(" | ".join(date_info))
                date_run = date_para.runs[0]
                date_run.italic = True
            
            # GPA if provided
            if edu.gpa:
                gpa_para = doc.add_paragraph(f"GPA: {edu.gpa}")
            
            # Relevant coursework or details
            if edu.coursework:
                for course in edu.coursework:
                    bullet_para = doc.add_paragraph(f"• {course}")
                    bullet_para.paragraph_format.left_indent = Inches(0.25)
    
    def _add_skills_section(self, doc: Document, skills: Skills) -> None:
        """
        Add skills section.
        
        Args:
            doc: Document object
            skills: Skills model instance
        """
        doc.add_heading("Skills", level=2)
        
        # Handle categorized skills
        if skills.categories:
            for skill_category in skills.categories:
                if skill_category.skills:
                    skills_text = ", ".join(skill_category.skills)
                    para = doc.add_paragraph()
                    
                    # Add category in bold
                    category_run = para.add_run(f"{skill_category.name}: ")
                    category_run.bold = True
                    
                    # Add skills
                    para.add_run(skills_text)
        
        # Handle raw skills list
        elif skills.raw_skills:
            skills_text = ", ".join(skills.raw_skills)
            para = doc.add_paragraph(skills_text)
    
    def _add_projects_section(self, doc: Document, projects: List[Project]) -> None:
        """
        Add projects section.
        
        Args:
            doc: Document object
            projects: List of project items
        """
        doc.add_heading("Projects", level=2)
        
        for project in projects:
            # Project name
            title_para = doc.add_heading(project.name, level=3)
            
            # Technologies
            if project.technologies:
                tech_text = ", ".join(project.technologies)
                tech_para = doc.add_paragraph()
                tech_run = tech_para.add_run(f"Technologies: ")
                tech_run.bold = True
                tech_para.add_run(tech_text)
            
            # Description
            if project.description:
                doc.add_paragraph(project.description)
            
            # URL
            if project.url:
                url_para = doc.add_paragraph(f"URL: {project.url}")
    
    def _add_certifications_section(self, doc: Document, certifications: List[Certification]) -> None:
        """
        Add certifications section.
        
        Args:
            doc: Document object
            certifications: List of certification items
        """
        doc.add_heading("Certifications", level=2)
        
        for cert in certifications:
            # Certification name and issuer
            cert_para = doc.add_paragraph()
            
            # Name in bold
            name_run = cert_para.add_run(cert.name)
            name_run.bold = True
            
            # Issuer
            if cert.issuer:
                cert_para.add_run(f" - {cert.issuer}")
            
            # Date
            if cert.date:
                cert_para.add_run(f" ({cert.date})")
            
            # URL (if available)
            if hasattr(cert, 'url') and cert.url:
                url_para = doc.add_paragraph(f"Verification: {cert.url}")
    
    def validate_output(self, docx_bytes: bytes) -> bool:
        """
        Validate generated DOCX for basic compliance.
        
        Args:
            docx_bytes: DOCX content to validate
            
        Returns:
            bool: True if validation passes
            
        Raises:
            ValueError: If validation fails
        """
        # Basic DOCX validation
        if not docx_bytes:
            raise ValueError("DOCX content is empty")
        
        # Check for DOCX file signature (ZIP format)
        if not docx_bytes.startswith(b'PK'):
            raise ValueError("Invalid DOCX format - missing ZIP signature")
        
        # Check minimum size
        if len(docx_bytes) < 5000:
            raise ValueError("DOCX file too small - possible generation error")
        
        logger.info("DOCX validation passed")
        return True
    
    @classmethod
    def from_config(cls, config: OutputConfig, template_name: str = "professional") -> "DOCXGenerator":
        """
        Create DOCXGenerator from OutputConfig with template selection.
        
        Args:
            config: Output configuration containing DOCX config
            template_name: Name of the template to use
            
        Returns:
            DOCXGenerator: Configured generator instance
        """
        return cls(config.docx, template_name)